from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report"
OUT_DIR.mkdir(exist_ok=True)


TITLE = "基于 Atlas 200I DK A2 的 RTC 驱动修改与验证实验报告"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9.5)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(12.0)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True)
        set_cell_text(cells[1], value)
        set_cell_shading(cells[0], "E8EEF5")
        for c in cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_arch_table(doc):
    rows = [
        ("硬件层", "RX8900/RV8803 兼容 RTC 芯片，挂在 I2C5 总线上，设备地址为 0x32。"),
        ("设备树层", '通过 compatible = "epson,rx8900" 和 reg = <0x32> 描述硬件连接关系。'),
        ("驱动层", "Linux 标准 RTC 驱动 rtc-rv8803.c 通过 I2C 访问芯片寄存器，并实现 read_time / set_time 等操作。"),
        ("内核子系统", "RTC core 统一管理 /dev/rtc0、/sys/class/rtc/rtc0 与 ioctl 调用。"),
        ("用户态应用", "rtc_check 打开 /dev/rtc0，并通过 ioctl(RTC_RD_TIME) 读取硬件时间。"),
        ("协作层", "使用 Git 管理补丁脚本、测试程序、运行脚本和实验记录，开发板可 clone 仓库复现实验。"),
    ]
    add_kv_table(doc, rows)


def add_version_table(doc):
    rows = [
        ("v1.0", "创建初始仓库，确定 RTC 驱动修改目标与 README。"),
        ("v1.1", "新增 rtc_check.c，通过 ioctl(RTC_RD_TIME) 读取 /dev/rtc0。"),
        ("v1.2", "新增 patch_rtc_rv8803.py，自动给 rtc-rv8803.c 增加可观察日志。"),
        ("v1.3", "新增开发板加载和运行脚本，固化 insmod、lsmod、dmesg、hwclock 验证流程。"),
        ("v1.4", "补充实验报告、Git 上传步骤和开发板拉取步骤。"),
        ("v1.5", "升级 rtc_check 为系统验证工具，支持 --device、--compare、--proc、--version。"),
    ]
    add_kv_table(doc, rows)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.color.rgb = RGBColor(31, 77, 120) if level > 1 else RGBColor(46, 116, 181)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for r in runs:
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(10.5)
    return p


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(31, 77, 120)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def build_docx():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(10)
    r = meta.add_run("嵌入式系统课程实验 | 方向：代码版本控制与协同开发 / Atlas RTC 驱动修改")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)

    add_heading(doc, "一、实验选题与总体判断", 1)
    add_para(
        doc,
        "本实验选择课程内容安排中的“代码版本控制与协同开发”方向，围绕 Atlas 200I DK A2 开发板上的 Linux RTC 驱动进行修改、编译、加载与用户态验证。最新版项目为 v1.5。实验对象位于真实 RTC 驱动链路中：设备树描述 I2C 总线上的 RTC 芯片，内核驱动接入 Linux RTC 子系统，用户程序通过 /dev/rtc0 读取硬件时间。",
    )
    add_para(
        doc,
        "本实验的重点在于用 Git 管理驱动修改过程，用脚本固化编译与运行流程，并通过开发板端程序证明修改确实进入了内核驱动路径。实验内容覆盖硬件描述、内核驱动、用户态访问和协同开发记录，能够体现嵌入式系统中软件与硬件之间的贯通关系。",
    )

    add_heading(doc, "二、系统链路分析：从硬件到用户态", 1)
    add_para(
        doc,
        "根据 RTC 驱动文档，Atlas 200I DK A2 的 RTC 设备为 Epson RX8900/RV8803 兼容芯片，设备树节点位于 I2C5 总线，地址为 0x32。Linux 中 RTC 并不是孤立驱动，而是由设备树、I2C 框架、RTC class、字符设备接口和用户态程序共同构成的系统链路。",
    )
    add_arch_table(doc)
    add_code_block(
        doc,
        '设备树关键片段：\n&i2c5_ao {\n    rtc@32 {\n        compatible = "epson,rx8900";\n        reg = <0x32>;\n    };\n};',
    )

    add_heading(doc, "三、驱动修改方案", 1)
    add_para(
        doc,
        "本实验不改变 RTC 时间读写算法本身，而是在真实驱动 rtc-rv8803.c 中加入可观察的日志点。这样既避免破坏硬件时序，又能证明驱动路径被实际触发。修改点包括：",
    )
    add_bullet(doc, "probe：I2C 设备与驱动匹配成功时输出 [atlas-rtc-demo] 日志。")
    add_bullet(doc, "read_time：用户态读取 /dev/rtc0 时输出 read_time called，验证 ioctl 调用进入具体驱动。")
    add_bullet(doc, "set_time：用户态或系统调用写入 RTC 时间时输出 set_time called，便于扩展验证。")
    add_para(
        doc,
        "驱动编译遵循文档中的 kernel 编译路径。RTC 属于 Linux 标准内核驱动，需在 menuconfig 中选择 Device Drivers -> Real Time Clock -> Micro Crystal RV8803, Epson RX8900，并建议编译为模块 CONFIG_RTC_DRV_RV8803=m。编译成功后，目标模块一般位于 output/kernel_modules/rtc-rv8803.ko。",
    )

    add_heading(doc, "四、Git 项目与开发板应用", 1)
    add_para(
        doc,
        "项目仓库设计为可复现实验流程，而不只是保存源码片段。仓库中包含驱动补丁脚本、编译提示脚本、开发板加载脚本、用户态测试程序和实验说明。开发板拉取仓库后，可以直接编译并运行测试程序。",
    )
    add_code_block(
        doc,
        "仓库核心结构：\n"
        "include/project_version.h           项目版本号，当前为 v1.5\n"
        "src/rtc_check.c                     用户态 RTC 读取程序\n"
        "scripts/patch_rtc_rv8803.py         自动给 rtc-rv8803.c 加日志\n"
        "scripts/build_kernel_module_on_host.sh  编译主机端补丁入口\n"
        "scripts/load_module_on_board.sh     开发板加载 rtc-rv8803.ko\n"
        "scripts/run_on_board.sh             编译并运行 rtc_check\n"
        "CHANGELOG.md                        版本迭代记录\n"
        "README.md / GIT_AND_BOARD_STEPS.md  实验说明与 Git 流程",
    )
    add_para(
        doc,
        "v1.5 版本将用户态程序 rtc_check 从单一读时间程序升级为开发板端验证工具。它支持 --device 指定 RTC 节点，--compare 对比 RTC 时间与系统时间，--proc 读取 /proc/driver/rtc 摘要，--version 输出项目版本。其核心调用仍然是打开 /dev/rtc0 并执行 ioctl(fd, RTC_RD_TIME, &rtc_tm)。这个调用会先进入 rtc-dev.c，再通过 rtc_class_ops 分发到 rtc-rv8803.c 的 read_time 函数。因此如果运行 rtc_check 后在 dmesg 中看到 [atlas-rtc-demo] read_time called，就说明驱动修改已经生效。",
    )

    add_heading(doc, "五、版本迭代与协同开发", 1)
    add_para(
        doc,
        "为了体现“代码版本控制与协同开发”，项目不是一次性提交成品，而是按照功能逐步迭代。每一版都对应一个清晰的工程目标：先建立仓库，再增加用户态测试，再加入驱动补丁脚本，随后补充开发板部署脚本，最后升级为可收集多层证据的系统验证工具。",
    )
    add_version_table(doc)

    add_heading(doc, "六、实验步骤与验证方法", 1)
    add_number(doc, "在 Git 平台创建仓库，提交 atlas-rtc-driver-demo 项目，形成可追踪的代码历史。")
    add_number(doc, "在编译主机上执行补丁脚本，修改 /opt/Ascend310B-source/drivers/rtc/rtc-rv8803.c。")
    add_number(doc, "执行 bash build.sh kernel，并在 menuconfig 中将 RTC_DRV_RV8803 选为模块。")
    add_number(doc, "将 output/kernel_modules/rtc-rv8803.ko 上传到开发板 /run 目录。")
    add_number(doc, "开发板 git clone 项目，运行 load_module_on_board.sh 加载模块，再运行 run_on_board.sh 读取 RTC 时间。")
    add_code_block(
        doc,
        "关键验收命令：\n"
        "git log --oneline -3\n"
        "ls -l /dev/rtc*\n"
        "lsmod | grep -E 'rtc|rv8803'\n"
        "./rtc_check --device /dev/rtc0 --compare --proc\n"
        "dmesg | grep atlas-rtc-demo\n"
        "cat /proc/driver/rtc\n"
        "hwclock -r -f /dev/rtc0",
    )

    add_heading(doc, "七、结果分析与课程意义", 1)
    add_para(
        doc,
        "本实验的验证不是只看程序是否能运行，而是看整个软硬件链路是否闭合：设备树负责描述硬件，I2C 框架负责发现和通信，RTC 驱动负责读写芯片时间寄存器，RTC 子系统向上提供统一接口，用户态程序通过 /dev/rtc0 触发驱动操作。新增日志提供了从用户态调用到内核驱动函数的证据链。",
    )
    add_para(
        doc,
        "从工程实践角度，Git 的作用不是形式化提交，而是保证修改可追踪、可协作、可回退。补丁脚本使驱动修改从“手工改文件”变成“可复现流程”；开发板端脚本则把加载模块、读取 RTC、检查 dmesg 固化为标准验证步骤。v1.5 新增的 --compare 和 --proc 进一步把用户态输出、内核 proc 信息和驱动日志放到同一次验收运行中，符合嵌入式开发中“主机交叉编译、目标板运行验证、日志定位问题”的基本工作流。",
    )
    add_para(
        doc,
        "本实验仍有可扩展空间：后续可以进一步实现或验证 RTC 闹钟中断、RTC_VL_READ 电压低标志、系统时间与 RTC 时间同步、掉电保持测试，以及对 I2C 读写错误进行更细粒度的日志分析。当前版本已经能够展示从硬件节点描述到用户态验证的完整链路，并保留了进一步扩展驱动功能的空间。",
    )

    add_heading(doc, "八、结论", 1)
    add_para(
        doc,
        "本实验基于 Atlas 200I DK A2 的 RTC 驱动完成了一个可管理、可编译、可部署、可验证、可迭代的小型嵌入式系统实践。实验将设备树、I2C 总线、Linux RTC 子系统、内核模块、用户态 ioctl 和 Git 协作流程联系起来，形成了从硬件描述到用户态验证的完整闭环。v1.5 版本在读取 RTC 时间的基础上增加了系统时间对比、/proc/driver/rtc 摘要读取和版本标识输出，使开发板端运行结果能够同时反映用户态接口、内核 RTC 子系统和驱动日志三个层面的状态。通过该过程，可以体现嵌入式系统实验中软硬件协同、驱动修改、目标板验证和版本化协作开发之间的关系。",
    )

    out = OUT_DIR / "Atlas_RTC驱动修改实验报告_v1.5_formal.docx"
    doc.save(out)
    return out


def build_markdown():
    md = f"""# {TITLE}

嵌入式系统课程实验 | 方向：代码版本控制与协同开发 / Atlas RTC 驱动修改

## 一、实验选题与总体判断

本实验选择课程内容安排中的“代码版本控制与协同开发”方向，围绕 Atlas 200I DK A2 开发板上的 Linux RTC 驱动进行修改、编译、加载与用户态验证。最新版项目为 v1.5。实验对象位于真实 RTC 驱动链路中：设备树描述 I2C 总线上的 RTC 芯片，内核驱动接入 Linux RTC 子系统，用户程序通过 `/dev/rtc0` 读取硬件时间。

本实验的重点在于用 Git 管理驱动修改过程，用脚本固化编译与运行流程，并通过开发板端程序证明修改确实进入了内核驱动路径。实验内容覆盖硬件描述、内核驱动、用户态访问和协同开发记录，能够体现嵌入式系统中软件与硬件之间的贯通关系。

## 二、系统链路分析

根据 RTC 驱动文档，Atlas 200I DK A2 的 RTC 设备为 Epson RX8900/RV8803 兼容芯片，设备树节点位于 I2C5 总线，地址为 `0x32`。Linux 中 RTC 并不是孤立驱动，而是由设备树、I2C 框架、RTC class、字符设备接口和用户态程序共同构成的系统链路。

设备树关键片段：

```dts
&i2c5_ao {{
    rtc@32 {{
        compatible = "epson,rx8900";
        reg = <0x32>;
    }};
}};
```

用户态程序调用链路为：

```text
rtc_check -> /dev/rtc0 -> ioctl(RTC_RD_TIME) -> rtc-dev.c
          -> rtc_class_ops.read_time -> rtc-rv8803.c -> I2C -> RTC 芯片
```

## 三、驱动修改方案

本实验不改变 RTC 时间读写算法本身，而是在真实驱动 `rtc-rv8803.c` 中加入可观察的日志点。这样既避免破坏硬件时序，又能证明驱动路径被实际触发。

- `probe`：I2C 设备与驱动匹配成功时输出 `[atlas-rtc-demo]` 日志。
- `read_time`：用户态读取 `/dev/rtc0` 时输出 `read_time called`。
- `set_time`：用户态或系统调用写入 RTC 时间时输出 `set_time called`。

驱动编译遵循文档中的 kernel 编译路径。RTC 属于 Linux 标准内核驱动，需在 menuconfig 中选择 `Device Drivers -> Real Time Clock -> Micro Crystal RV8803, Epson RX8900`，并建议编译为模块 `CONFIG_RTC_DRV_RV8803=m`。编译成功后，目标模块一般位于 `output/kernel_modules/rtc-rv8803.ko`。

## 四、Git 项目与开发板应用

项目仓库设计为可复现实验流程，而不只是保存源码片段。仓库中包含驱动补丁脚本、编译提示脚本、开发板加载脚本、用户态测试程序和实验说明。开发板拉取仓库后，可以直接编译并运行测试程序。

核心结构：

```text
include/project_version.h
src/rtc_check.c
scripts/patch_rtc_rv8803.py
scripts/build_kernel_module_on_host.sh
scripts/load_module_on_board.sh
scripts/run_on_board.sh
CHANGELOG.md
README.md / GIT_AND_BOARD_STEPS.md
```

v1.5 版本将用户态程序 `rtc_check` 从单一读时间程序升级为开发板端验证工具。它支持 `--device` 指定 RTC 节点，`--compare` 对比 RTC 时间与系统时间，`--proc` 读取 `/proc/driver/rtc` 摘要，`--version` 输出项目版本。其核心调用仍然是打开 `/dev/rtc0` 并执行 `ioctl(fd, RTC_RD_TIME, &rtc_tm)`。这个调用会先进入 `rtc-dev.c`，再通过 `rtc_class_ops` 分发到 `rtc-rv8803.c` 的 `read_time` 函数。因此如果运行 `rtc_check` 后在 `dmesg` 中看到 `[atlas-rtc-demo] read_time called`，就说明驱动修改已经生效。

## 五、版本迭代与协同开发

为了体现“代码版本控制与协同开发”，项目不是一次性提交成品，而是按照功能逐步迭代。每一版都对应一个清晰的工程目标：先建立仓库，再增加用户态测试，再加入驱动补丁脚本，随后补充开发板部署脚本，最后升级为可收集多层证据的系统验证工具。

| 版本 | 迭代内容 |
|---|---|
| v1.0 | 创建初始仓库，确定 RTC 驱动修改目标与 README。 |
| v1.1 | 新增 `rtc_check.c`，通过 `ioctl(RTC_RD_TIME)` 读取 `/dev/rtc0`。 |
| v1.2 | 新增 `patch_rtc_rv8803.py`，自动给 `rtc-rv8803.c` 增加可观察日志。 |
| v1.3 | 新增开发板加载和运行脚本，固化 `insmod`、`lsmod`、`dmesg`、`hwclock` 验证流程。 |
| v1.4 | 补充实验报告、Git 上传步骤和开发板拉取步骤。 |
| v1.5 | 升级 `rtc_check` 为系统验证工具，支持 `--device`、`--compare`、`--proc`、`--version`。 |

## 六、实验步骤与验证方法

1. 在 Git 平台创建仓库，提交 `atlas-rtc-driver-demo` 项目。
2. 在编译主机上执行补丁脚本，修改 `/opt/Ascend310B-source/drivers/rtc/rtc-rv8803.c`。
3. 执行 `bash build.sh kernel`，并在 menuconfig 中将 `RTC_DRV_RV8803` 选为模块。
4. 将 `output/kernel_modules/rtc-rv8803.ko` 上传到开发板 `/run` 目录。
5. 开发板 `git clone` 项目，运行 `load_module_on_board.sh` 加载模块，再运行 `run_on_board.sh` 读取 RTC 时间。

关键验收命令：

```bash
git log --oneline -3
ls -l /dev/rtc*
lsmod | grep -E 'rtc|rv8803'
./rtc_check --device /dev/rtc0 --compare --proc
dmesg | grep atlas-rtc-demo
cat /proc/driver/rtc
hwclock -r -f /dev/rtc0
```

## 七、结果分析与课程意义

本实验的验证不是只看程序是否能运行，而是看整个软硬件链路是否闭合：设备树负责描述硬件，I2C 框架负责发现和通信，RTC 驱动负责读写芯片时间寄存器，RTC 子系统向上提供统一接口，用户态程序通过 `/dev/rtc0` 触发驱动操作。新增日志提供了从用户态调用到内核驱动函数的证据链。

从工程实践角度，Git 的作用不是形式化提交，而是保证修改可追踪、可协作、可回退。补丁脚本使驱动修改从“手工改文件”变成“可复现流程”；开发板端脚本则把加载模块、读取 RTC、检查 `dmesg` 固化为标准验证步骤。v1.5 新增的 `--compare` 和 `--proc` 进一步把用户态输出、内核 proc 信息和驱动日志放到同一次验收运行中，符合嵌入式开发中“主机交叉编译、目标板运行验证、日志定位问题”的基本工作流。

本实验仍有可扩展空间：后续可以进一步实现或验证 RTC 闹钟中断、`RTC_VL_READ` 电压低标志、系统时间与 RTC 时间同步、掉电保持测试，以及对 I2C 读写错误进行更细粒度的日志分析。当前版本已经能够展示从硬件节点描述到用户态验证的完整链路，并保留了进一步扩展驱动功能的空间。

## 八、结论

本实验基于 Atlas 200I DK A2 的 RTC 驱动完成了一个可管理、可编译、可部署、可验证、可迭代的小型嵌入式系统实践。实验将设备树、I2C 总线、Linux RTC 子系统、内核模块、用户态 ioctl 和 Git 协作流程联系起来，形成了从硬件描述到用户态验证的完整闭环。v1.5 版本在读取 RTC 时间的基础上增加了系统时间对比、`/proc/driver/rtc` 摘要读取和版本标识输出，使开发板端运行结果能够同时反映用户态接口、内核 RTC 子系统和驱动日志三个层面的状态。通过该过程，可以体现嵌入式系统实验中软硬件协同、驱动修改、目标板验证和版本化协作开发之间的关系。
"""
    out = OUT_DIR / "Atlas_RTC驱动修改实验报告_v1.5_formal.md"
    out.write_text(md, encoding="utf-8")
    return out


if __name__ == "__main__":
    print(build_docx())
    print(build_markdown())
