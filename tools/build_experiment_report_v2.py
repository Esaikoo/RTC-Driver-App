from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report"
OUT_DIR.mkdir(exist_ok=True)

TITLE = "基于 Atlas 200I DK A2 的 RTC 驱动修改与协同开发实验报告"
OUT_FILE = OUT_DIR / "Atlas_RTC驱动协同开发实验报告_v2.0.docx"


def set_east_asia(run, font="微软雅黑"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    set_east_asia(r)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text, size=10.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.font.size = Pt(size)
    set_east_asia(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        set_east_asia(r)
        r.font.color.rgb = RGBColor(46, 116, 181) if level == 1 else RGBColor(31, 77, 120)
        r.font.size = Pt(13.5 if level == 1 else 11.2)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.size = Pt(9.8)
    set_east_asia(r)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(8.8)
    r.font.color.rgb = RGBColor(50, 50, 50)


def add_kv_table(doc, rows, widths=(3.0, 13.0)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(widths[0])
        cells[1].width = Cm(widths[1])
        set_cell_text(cells[0], key, bold=True)
        set_cell_text(cells[1], value)
        set_cell_shading(cells[0], "E8EEF5")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08


def add_footer(doc):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("RTC Driver App v2.0 | Atlas 200I DK A2")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 120)
    set_east_asia(run)


def build():
    doc = Document()
    configure_doc(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(0, 51, 102)
    set_east_asia(r)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(7)
    r = meta.add_run("课程方向：代码版本控制与协同开发 | 项目版本：v2.0 | 仓库：Esaikoo/RTC-Driver-App")
    r.font.size = Pt(9.2)
    r.font.color.rgb = RGBColor(90, 90, 90)
    set_east_asia(r)

    add_heading(doc, "一、实验目标与选题定位")
    add_para(
        doc,
        "本实验选择课程第 6 项中的“代码版本控制与协同开发”方向，以 Atlas 200I DK A2 开发板上的 RTC 驱动为对象，在原有 v1.5 仓库基础上升级到 v2.0。实验不是只做一个用户态小程序，而是把硬件 RTC、设备树、Linux RTC 子系统、具体 I2C 驱动、用户态验收工具和 Git 协作流程串成一条可复现的系统链路。",
    )
    add_para(
        doc,
        "v2.0 的目标是让验收更接近工程实践：连续采样 RTC 时间，对比系统时间，设置漂移阈值，并输出 JSON 证据，便于不同同学在不同开发板或不同分支上复现实验结果。",
    )

    add_heading(doc, "二、系统链路：软硬件结合关系")
    add_kv_table(
        doc,
        [
            ("硬件层", "RTC 芯片为 RX8900/RV8803 兼容器件，挂接在 Atlas 开发板 I2C 总线上，硬件地址由设备树 reg=<0x32> 描述。"),
            ("驱动层", "rtc-rv8803.c 通过 I2C 框架读写 RTC 寄存器，并向 Linux RTC core 注册 read_time、set_time 等操作。"),
            ("内核接口", "RTC core 统一暴露 /dev/rtc0、/proc/driver/rtc 和 hwclock 可用接口，屏蔽底层芯片差异。"),
            ("用户态", "rtc_check 打开 /dev/rtc0，通过 ioctl(RTC_RD_TIME) 触发内核 RTC 读路径，并与系统时间进行校验。"),
            ("协作层", "Git 记录从 v1.0 到 v2.0 的功能演进，同学可 clone、建分支、提交验收证据并互相 review。"),
        ],
    )
    add_code(
        doc,
        '设备树关键关系：compatible = "epson,rx8900"; reg = <0x32>;\n用户态调用链：rtc_check -> /dev/rtc0 -> RTC core -> rtc-rv8803.c -> I2C -> RTC 芯片',
    )

    add_heading(doc, "三、版本控制与 v2.0 改进")
    add_para(
        doc,
        "原仓库已经包含 v1.0 至 v1.5 的提交历史：从初始 README、用户态读取 RTC、驱动 patch 脚本、开发板运行脚本，到 v1.5 的系统验收工具。v2.0 在此基础上继续演进，体现协作开发中“从可运行到可验收、从人工观察到证据沉淀”的改进过程。",
    )
    add_kv_table(
        doc,
        [
            ("连续采样", "新增 --samples 和 --interval，可在一次验收中读取多次 RTC，避免单点读数偶然性。"),
            ("漂移阈值", "新增 --max-drift，当 RTC 与系统时间误差超过阈值时返回非零退出码，便于脚本化验收。"),
            ("JSON 证据", "新增 --json 输出 JSON Lines，保存为 rtc_evidence_v2.jsonl 后可随报告或分支提交。"),
            ("验收脚本", "run_on_board.sh 同时收集人工可读输出、JSON 输出、hwclock 与 dmesg 日志。"),
        ],
    )
    add_code(
        doc,
        "sudo ./rtc_check --device /dev/rtc0 --compare --max-drift 300 --samples 3 --interval 1 --json | tee rtc_evidence_v2.jsonl",
    )

    add_heading(doc, "四、关键实现说明")
    add_bullet(doc, "参数解析：在原 --device、--compare、--proc 基础上增加 --samples、--interval、--max-drift、--json，错误参数返回 2。")
    add_bullet(doc, "时间转换：将 struct rtc_time 转换为 time_t 后与 time(NULL) 比较，输出 drift_seconds。")
    add_bullet(doc, "脚本验收：若绝对漂移超过阈值，程序返回 3，使自动脚本能直接判断失败。")
    add_bullet(doc, "证据格式：JSON 中包含 sample、version、device、rtc_time、rtc_epoch、system_epoch、drift_seconds、status。")

    add_heading(doc, "五、协同开发过程")
    add_para(
        doc,
        "协同开发的重点不是多人同时改同一个文件，而是把需求拆成可审阅、可回退、可验证的版本演进。本项目中，一名同学维护初始仓库，其他同学通过 clone 仓库、创建功能分支、提交修改、运行开发板验收命令、再合并回主分支的方式完成协作。v2.0 的修改适合作为一次功能分支提交，因为它同时改变了用户态工具、运行脚本、README、CHANGELOG 和验收说明，能够清楚展示一次完整工程迭代。",
    )
    add_kv_table(
        doc,
        [
            ("需求提出", "v1.5 已能读取 RTC，但验收证据主要依赖人工观察，缺少连续采样和机器可读记录。"),
            ("分支开发", "在 feature/rtc-v2-evidence 分支中实现多次采样、漂移阈值和 JSON 输出。"),
            ("本地审查", "通过 git diff 检查代码、脚本和文档是否围绕同一目标变化，避免只改版本号。"),
            ("板端验收", "在 Atlas 上运行 run_on_board.sh，保存终端输出、dmesg 片段和 rtc_evidence_v2.jsonl。"),
            ("合并归档", "将报告、验收记录和代码提交合并到主分支，形成可追溯的课程实验版本。"),
        ],
    )

    add_heading(doc, "六、验收记录设计")
    add_para(
        doc,
        "v2.0 把验收分为四类证据：设备节点证据、用户态读取证据、内核路径证据和可复核数据证据。这样做的好处是，某一层出现问题时可以定位：如果 /dev/rtc0 不存在，先查设备树和驱动注册；如果 ioctl 失败，查权限或 RTC core；如果 dmesg 没有标记，说明自定义模块可能没有替换成功；如果漂移超阈值，则需要检查系统时间同步或 RTC 电池/初始化状态。",
    )
    add_kv_table(
        doc,
        [
            ("设备节点", "ls -l /dev/rtc*，确认 RTC 字符设备已注册。"),
            ("用户态读取", "./rtc_check --device /dev/rtc0 --compare --max-drift 300 --samples 3。"),
            ("内核日志", "dmesg | grep atlas-rtc-demo，确认 read_time/probe 日志来自修改后的驱动。"),
            ("系统工具", "hwclock -r -f /dev/rtc0，与自写程序互相印证。"),
            ("数据归档", "JSON Lines 每行对应一次样本，适合提交到 Git 或粘贴到报告附录。"),
        ],
    )
    add_code(
        doc,
        'JSON 样例：\n{"sample":1,"version":"v2.0","device":"/dev/rtc0","rtc_time":"2026-06-17 02:10:30","drift_seconds":1,"status":"pass"}',
    )

    add_heading(doc, "七、开发板验收流程")
    add_para(doc, "在 Atlas 200I DK A2 上通过 Type-C/RNDIS 或网口登录后，执行以下流程即可复现实验：")
    add_code(
        doc,
        "git clone https://github.com/Esaikoo/RTC-Driver-App.git\ncd RTC-Driver-App\nmake\nsudo bash scripts/run_on_board.sh /dev/rtc0\ndmesg | tail -80 | grep -iE 'rtc|rv8803|atlas-rtc-demo'",
    )
    add_para(
        doc,
        "如果已替换带日志的 rtc-rv8803.ko，则 dmesg 中应能看到 [atlas-rtc-demo] 标记；如果只验证用户态读取，也应看到 /dev/rtc0 存在、RTC 时间可读、JSON 状态为 pass。这说明用户态程序、内核 RTC 子系统和硬件 RTC 之间形成了闭环。",
    )

    doc.add_page_break()
    add_heading(doc, "八、问题分析与改进方向")
    add_para(
        doc,
        "实验中需要注意版本一致性。课程 PPT 特别提醒 Atlas 开发板 prebuilt 镜像的驱动版本可能与下载的 SDK 版本不同，如果内核、驱动和固件版本不一致，轻则模块无法加载，重则开发板重启。因此，驱动编译前应确认内核源码、交叉编译工具链、板端 uname -a 和驱动包版本一致。v2.0 工具本身不替代驱动编译，但它提供了一个清晰的验收入口，可以在驱动替换后立即验证修改是否生效。",
    )
    add_bullet(doc, "后续可把 JSON 证据上传到 report/evidence 目录，并按日期命名，形成多名同学的实验记录。")
    add_bullet(doc, "可增加 GitHub Actions 的基础语法检查，但真正的 RTC 验证仍必须在 Atlas 硬件上完成。")
    add_bullet(doc, "可扩展 set_time 验证，但要谨慎避免错误写入影响系统时间和其他实验。")

    add_para(
        doc,
        "从工程角度看，RTC 驱动实验还可以继续做三类增强。第一类是可靠性增强，例如在用户态连续读取时检查秒值是否单调变化，发现时间停滞时给出告警；第二类是协作增强，例如要求每名组员在独立分支提交一次板端验收记录，由组长统一 review 后合并；第三类是报告增强，例如把 dmesg、lsmod、hwclock 和 JSON 证据对应到系统链路图中的不同层级，使验收材料不仅能证明“运行过”，还可以说明“为什么能运行”。",
    )
    add_para(
        doc,
        "需要避免的做法是只修改 README 或只截取终端登录图。课程要求强调“系统”和“软硬件结合”，因此报告中必须把硬件连接、内核驱动、用户态接口和 Git 协作联系起来。v2.0 的代码变化正是围绕这个目标设计：用户态工具提供可重复证据，脚本统一板端流程，文档说明版本演进，最后由 Git 保存每次修改。这样，版本控制不再只是提交作业的手段，而成为实验复现和团队协作的一部分。",
    )

    doc.add_page_break()
    add_heading(doc, "九、实验认识")
    add_para(
        doc,
        "通过本实验可以看到，嵌入式系统开发不是单纯写 C 程序，也不是孤立修改驱动。设备树负责把硬件事实告诉内核，驱动把芯片寄存器转换成标准 RTC 操作，用户态程序通过统一接口完成验证，而 Git 让修改、验收和协作过程可追踪。v2.0 的改进把一次性演示升级为可重复、可比较、可审阅的工程流程，较好体现了课程要求中的“系统”和“软硬件结合”。",
    )
    add_para(
        doc,
        "对“系统”的理解也在实验中变得具体。RTC 芯片本身只负责保持时间，但要让这个硬件能力被应用程序使用，需要设备树描述硬件连接，需要 I2C 驱动完成总线通信，需要 RTC 子系统抽象标准接口，需要字符设备把能力暴露给用户态，还需要脚本和文档把操作流程固定下来。任何一层缺失，最终的 ./rtc_check 输出都不会可靠。",
    )
    add_para(
        doc,
        "对“软硬件结合”的理解则体现在验证方法上。只看 C 代码，无法证明硬件真的被访问；只看开发板登录截图，也无法说明调用经过了哪条内核路径。因此本实验同时保留用户态输出、/proc/driver/rtc、hwclock、dmesg 和 JSON 证据，用多种来源互相印证。报告中的每条证据都能对应到系统链路中的一层，这比单纯展示运行结果更能说明理论和实践的贯通。",
    )
    add_para(
        doc,
        "最后，Git 协作让实验具备可复现性。v2.0 的一次提交不仅包含代码，还包含运行脚本、验收说明和报告生成脚本。后续同学可以基于同一仓库继续添加证据或改进驱动，而不是重新整理零散文件。这种“代码、文档、证据同步演进”的方式，正是嵌入式团队开发中非常重要的工程习惯。",
    )

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build()
